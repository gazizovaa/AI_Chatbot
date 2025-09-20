import os
import shutil
import re
import uuid
import json
import logging
from datetime import datetime
from typing import List, Optional, Dict, Any

# Flask imports
from flask import Flask, render_template, request, jsonify, redirect, url_for
from werkzeug.utils import secure_filename
from flask import send_from_directory

# Data processing imports
import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Try importing optional dependencies
try:
    import PyPDF2

    HAVE_PDF = True
except ImportError:
    HAVE_PDF = False

try:
    import docx

    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

try:
    import markdown

    HAVE_MARKDOWN = True
except ImportError:
    HAVE_MARKDOWN = False

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document

    HAVE_LANGCHAIN = True
except ImportError:
    HAVE_LANGCHAIN = False

try:
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import Chroma

    HAVE_LC_EMBEDDINGS = True
    HAVE_CHROMADB = True
except ImportError:
    HAVE_LC_EMBEDDINGS = False
    HAVE_CHROMADB = False

# Google API imports
try:
    import google.generativeai as genai

    HAVE_GEMINI = True
except ImportError:
    HAVE_GEMINI = False

# Load environment variables
from dotenv import load_dotenv

load_dotenv()

# --- Configuration ---
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
NEW_INDEX_FOLDER = 'vectorstore'
OLD_INDEX_FOLDER = 'old_vectorstore'

# Köhnə qovluğu silmək (əgər lazımsa)
if os.path.exists(OLD_INDEX_FOLDER) and os.path.isdir(OLD_INDEX_FOLDER):
    try:
        shutil.rmtree(OLD_INDEX_FOLDER)
        logging.info(f"Köhnə '{OLD_INDEX_FOLDER}' qovluğu uğurla silindi.")
    except Exception as e:
        logging.error(f"Köhnə '{OLD_INDEX_FOLDER}' qovluğu silinərkən xəta baş verdi: {e}")

# Yeni qovluqları yaratmaq
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(NEW_INDEX_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['INDEX_FOLDER'] = NEW_INDEX_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB limit

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md'}

# --- Logging ---
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- AI and Vector Store Setup ---
gemini_ai = None
if HAVE_GEMINI:
    GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
    if not GEMINI_API_KEY:
        logger.error("GEMINI_API_KEY is not set in environment variables.")
        HAVE_GEMINI = False
    else:
        try:
            genai.configure(api_key=GEMINI_API_KEY)
            gemini_ai = genai.GenerativeModel('gemini-1.5-flash')
            logger.info("Gemini AI model configured.")
        except Exception as e:
            logger.error(f"Gemini AI configuration error: {e}")
            HAVE_GEMINI = False
            gemini_ai = None
else:
    logger.warning("Google Generative AI library not found. AI features will be disabled.")

vector_store = None
if HAVE_LC_EMBEDDINGS and HAVE_CHROMADB:
    try:
        embeddings = HuggingFaceEmbeddings(
            model_name="intfloat/multilingual-e5-large",
            model_kwargs={'device': 'cpu'}
        )
        vector_store = Chroma(
            embedding_function=embeddings,
            persist_directory=app.config['INDEX_FOLDER']
        )
        logger.info("Vector store (ChromaDB) and embeddings configured.")
    except Exception as e:
        logger.error(f"Vector store configuration error: {e}")
        vector_store = None
else:
    logger.warning("LangChain embeddings or ChromaDB not found. Document search features will be disabled.")


# --- Helpers ---
def allowed_file(filename):
    """Checks if the file extension is allowed."""
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_document(filepath):
    """Parses text from various document types."""
    content = ""
    file_ext = filepath.rsplit('.', 1)[1].lower()

    try:
        if file_ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        elif file_ext == 'pdf' and HAVE_PDF:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    content += page.extract_text()
        elif file_ext == 'docx' and HAVE_DOCX:
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                content += para.text + "\n"
        elif file_ext == 'md' and HAVE_MARKDOWN:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
    except Exception as e:
        logger.error(f"Error parsing document {filepath}: {e}")
        return ""

    return content.strip()


def create_and_add_to_vector_store(doc_id: str, content: str, filename: str):
    """Splits text and adds it to the vector store."""
    if not HAVE_LANGCHAIN or not vector_store:
        logger.warning("LangChain or vector store not available")
        return False

    if not content.strip():
        logger.warning(f"Empty content for document {filename}")
        return False

    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            is_separator_regex=False
        )
        chunks = text_splitter.split_text(content)

        if not chunks:
            logger.warning(f"No chunks generated for document {filename}")
            return False

        docs = [
            Document(
                page_content=chunk,
                metadata={
                    "doc_id": doc_id,
                    "filename": filename,
                    "created_at": datetime.now().isoformat()
                }
            ) for chunk in chunks
        ]

        vector_store.add_documents(docs)
        logger.info(f"Successfully added {len(docs)} chunks for document {filename} to the vector store.")
        return True

    except Exception as e:
        logger.error(f"Failed to add document to vector store: {e}")
        return False


# --- Routes ---
@app.route('/')
def home():
    """Renders the main chatbot page."""
    return render_template('index.html')


@app.route('/documents')
def documents_page():
    """Renders the document management page."""
    return render_template('documents.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    """Handles document uploads."""
    if 'file' not in request.files:
        return jsonify({'error': 'Fayl tapılmadı'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'Fayl adı boşdur'}), 400

    if not file or not allowed_file(file.filename):
        return jsonify({
            'error': 'İcazə verilməyən fayl formatı. Yalnız .txt, .pdf, .docx və .md faylları qəbul olunur.'
        }), 400

    try:
        filename = secure_filename(file.filename)
        doc_id = str(uuid.uuid4())
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{doc_id}_{filename}")
        file.save(filepath)

        # Process the document for RAG
        content = parse_document(filepath)
        if not content:
            os.remove(filepath)
            return jsonify({
                'error': 'Faylın məzmunu oxuna bilmədi və ya fayl boşdur.'
            }), 400

        # Try to add to vector store
        success = create_and_add_to_vector_store(doc_id, content, filename)

        if success:
            return jsonify({
                'message': f"Fayl '{filename}' uğurla yükləndi və işləndi.",
                'doc_id': doc_id,
                'filename': filename
            }), 200
        else:
            # Keep the file even if vector store failed
            return jsonify({
                'message': f"Fayl '{filename}' yükləndi, lakin vektor bazaya əlavə edilmədi.",
                'doc_id': doc_id,
                'filename': filename
            }), 200

    except Exception as e:
        logger.error(f"File upload error: {str(e)}")
        return jsonify({
            'error': f"Fayl yüklənərkən xəta baş verdi: {str(e)}"
        }), 500


@app.route('/documents/list', methods=['GET'])
def list_documents():
    """Lists all processed documents."""
    try:
        if not vector_store:
            return get_files_from_uploads_folder()

        try:
            collection_count = vector_store._collection.count()
            if collection_count == 0:
                return get_files_from_uploads_folder()

            docs = vector_store.get(include=["metadatas"])['metadatas']

            # Get unique documents based on doc_id
            unique_docs = {doc['doc_id']: doc for doc in docs}.values()

            # Format for display
            doc_list = sorted([
                {
                    'doc_id': doc['doc_id'],
                    'filename': doc['filename'],
                    'created_at': doc['created_at']
                }
                for doc in unique_docs
            ], key=lambda x: x['created_at'], reverse=True)

            return jsonify(doc_list)
        except Exception as e:
            logger.error(f"Error accessing vector store: {e}")
            return get_files_from_uploads_folder()

    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        return get_files_from_uploads_folder()


def get_files_from_uploads_folder():
    """Uploads qovluğundakı faylları siyahıya alır."""
    try:
        files = []
        upload_folder = app.config['UPLOAD_FOLDER']

        if not os.path.exists(upload_folder):
            return jsonify([])

        for filename in os.listdir(upload_folder):
            if filename.endswith(('.txt', '.pdf', '.docx', '.md')):
                filepath = os.path.join(upload_folder, filename)
                created_time = os.path.getctime(filepath)

                # Fayl adından doc_id-ni çıxar
                parts = filename.split('_', 1)
                doc_id = parts[0] if len(parts) > 1 else 'unknown'
                original_filename = parts[1] if len(parts) > 1 else filename

                files.append({
                    'doc_id': doc_id,
                    'filename': original_filename,
                    'created_at': datetime.fromtimestamp(created_time).isoformat()
                })

        return jsonify(sorted(files, key=lambda x: x['created_at'], reverse=True))
    except Exception as e:
        logger.error(f"Error getting files from uploads folder: {str(e)}")
        return jsonify([])


@app.route('/documents/<doc_id>/delete', methods=['POST'])
def delete_document(doc_id):
    """Deletes a document and its chunks from the vector store."""
    try:
        filename = "Unknown"

        # Try to delete from vector store first
        if vector_store:
            try:
                existing_docs = vector_store.get(where={"doc_id": doc_id}, include=["metadatas"])

                if existing_docs and existing_docs.get('ids'):
                    filename = existing_docs['metadatas'][0]['filename']
                    vector_store.delete(where={"doc_id": doc_id})
                    logger.info(f"Successfully deleted all chunks for doc_id: {doc_id} from vector store.")
            except Exception as e:
                logger.error(f"Error deleting from vector store: {e}")

        # Delete the original file from the uploads folder
        deleted_file_count = 0
        upload_folder = app.config['UPLOAD_FOLDER']

        if os.path.exists(upload_folder):
            for f in os.listdir(upload_folder):
                if f.startswith(doc_id):
                    filepath_to_delete = os.path.join(upload_folder, f)
                    os.remove(filepath_to_delete)
                    logger.info(f"Successfully deleted original file: {filepath_to_delete}")
                    deleted_file_count += 1

                    # Get filename from file if not found in vector store
                    if filename == "Unknown":
                        parts = f.split('_', 1)
                        filename = parts[1] if len(parts) > 1 else f
                    break

        if deleted_file_count == 0:
            logger.warning(f"Original file for doc_id: {doc_id} not found.")
            return jsonify({'error': 'Sənəd tapılmadı.'}), 404

        return jsonify({'message': f"Sənəd '{filename}' uğurla silindi."}), 200

    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        return jsonify({'error': 'Sənəd silinərkən xəta baş verdi.'}), 500


@app.route('/chat', methods=['POST'])
def chat():
    """Handles user chat requests."""
    if not request.json or 'message' not in request.json:
        return jsonify({'error': 'Mesaj tapılmadı.'}), 400

    if not HAVE_GEMINI or not gemini_ai:
        return jsonify({'error': 'AI xidməti aktiv deyil. GEMINI_API_KEY-i yoxlayın.'}), 503

    user_message = request.json['message'].strip()

    if not user_message:
        return jsonify({'error': 'Boş mesaj göndərilə bilməz.'}), 400

    try:
        relevant_context = []
        sources = []

        # Search for relevant documents if vector store is available
        if vector_store:
            try:
                results = vector_store.similarity_search_with_score(user_message, k=5)
                relevant_context = [
                    {'content': doc.page_content, **doc.metadata}
                    for doc, score in results if score < 0.7  # Lower score means higher similarity
                ]

                # Extract sources from relevant documents
                sources = [
                    {'filename': ctx['filename'], 'doc_id': ctx['doc_id']}
                    for ctx in relevant_context
                ]
            except Exception as e:
                logger.error(f"Vector search error: {e}")

        # Prepare prompt based on available context
        if relevant_context:
            context_text = "\n\n".join([
                f"Sənəd: {ctx['filename']}\nMəzmun: {ctx['content']}"
                for ctx in relevant_context[:3]  # Limit to top 3 results
            ])

            prompt = f"""Sənədlərdə olan məlumatlara əsasən cavab ver. Azərbaycan dilində cavab ver.

Kontekst:
{context_text}

Sual: {user_message}

Cavab yalnız verilən kontekstdəki məlumatlara əsaslanmalıdır. Əgər kontekstdə cavab yoxdursa, bunu açıq şəkildə bildirin."""
        else:
            prompt = f"""Salam! Mən sizin şəxsi köməkçinizəm. 

Sualınız: {user_message}

Hazırda yüklənmiş sənədlərdə bu suala aid spesifik məlumat tapılmadı. Lakin ümumi məlumatlarımla sizə kömək edə bilərəm. 

Əgər daha dəqiq cavab almaq istəyirsinizsə, əlavə sənədlər yükləyin və ya sualınızı dəqiqləşdirin."""

        # Generate response using Gemini
        response = gemini_ai.generate_content(prompt)
        response_text = response.text if response and response.text else "Təəssüf ki, cavab hazırlana bilmədi."

        # Remove duplicates from sources
        unique_sources = []
        seen_docs = set()
        for source in sources:
            if source['doc_id'] not in seen_docs:
                unique_sources.append(source)
                seen_docs.add(source['doc_id'])

        return jsonify({
            'response': response_text,
            'sources': unique_sources[:3]  # Limit to 3 sources
        })

    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        return jsonify({
            'error': f'Cavab hazırlanarkən xəta baş verdi. Xahiş edirəm yenidən cəhd edin.'
        }), 500


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint."""
    status = {
        'status': 'OK',
        'gemini_ai': HAVE_GEMINI and gemini_ai is not None,
        'vector_store': vector_store is not None,
        'pdf_support': HAVE_PDF,
        'docx_support': HAVE_DOCX,
        'markdown_support': HAVE_MARKDOWN,
        'langchain_support': HAVE_LANGCHAIN
    }
    return jsonify(status)


if __name__ == '__main__':
    logger.info("Starting Flask application...")
    logger.info(f"Upload folder: {UPLOAD_FOLDER}")
    logger.info(f"Vector store folder: {NEW_INDEX_FOLDER}")
    app.run(debug=True, host='0.0.0.0', port=5000)